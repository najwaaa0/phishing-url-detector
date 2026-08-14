import os
import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

SECTIONS = [
    ("Abstract", "This report summarizes the Phishing URL Detector project (placeholder text)."),
    ("Introduction", "Short project description..."),
    # ...keep minimal; full content can be the one you previously used...
]

def add_section(doc, heading, text):
    doc.add_heading(heading, level=2)
    for para in text.split("\n\n"):
        doc.add_paragraph(para)

def add_figure(doc, title, img_path, width_inches=5.0):
    if not os.path.exists(img_path):
        return False
    doc.add_heading(title, level=3)
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return True

def main(out_path: str, figs_dir: str = None):
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    doc.add_heading("Phishing URL Detector — Final Report", level=0)
    doc.add_paragraph("Author: Project Team")
    doc.add_paragraph("Date: ")
    doc.add_page_break()

    for heading, text in SECTIONS:
        add_section(doc, heading, text)
    doc.add_page_break()

    if figs_dir:
        figs = [
            ("Class distribution", "class_distribution.png"),
            ("URL length distribution", "url_length_hist.png"),
            ("Feature correlation", "feature_corr.png"),
            ("Confusion matrix", "confusion_matrix.png"),
            ("ROC curve", "roc_curve.png"),
            ("Precision-Recall curve", "pr_curve.png"),
            ("Metrics vs threshold", "threshold_metrics.png"),
            ("Feature importance", "feature_importance.png"),
        ]
        for title, fname in figs:
            img_path = os.path.join(figs_dir, fname)
            add_figure(doc, title, img_path)

    doc.save(out_path)
    print(f"Saved DOCX report to: {out_path}")

if __name__ == '__main__':
    p = argparse.ArgumentParser()
    p.add_argument('--out', default='reports/FINAL_REPORT_with_figs.docx')
    p.add_argument('--figs-dir', default='reports/figs')
    args = p.parse_args()
    main(args.out, args.figs_dir)